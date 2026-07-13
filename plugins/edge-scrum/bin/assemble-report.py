#!/usr/bin/env python3
"""Assemble release planning report from checks.json and recommendations.json.

Produces both .md and .docx output. All tables are rendered from structured
data — no markdown parsing needed for the DOCX path.
"""

import argparse
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(__file__))
from _jira_transforms import load_json

JIRA_BASE = "https://redhat.atlassian.net/browse"
JIRA_KEY_RE = re.compile(r"(?<!\[)(?<!/)\b(OCPSTRAT-\d+|OCPEDGE-\d+|USHIFT-\d+|OCPBUGS-\d+)\b(?!\])")


def jira_linkify(text):
    """Replace bare Jira keys with markdown links."""
    return JIRA_KEY_RE.sub(rf"[\1]({JIRA_BASE}/\1)", text)


# --- Markdown Rendering ---


def render_md_table(headers, rows):
    lines = []
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join("---" for _ in headers) + " |")
    for row in rows:
        cells = [str(c).replace("|", "\\|") for c in row]
        while len(cells) < len(headers):
            cells.append("")
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def render_markdown(checks, recs, template, params):
    """Build the full markdown report."""
    meta = checks["meta"]

    header_map = {
        "VERSION": params["version"],
        "TODAY": params["today"],
        "FIRST": params["first_sprint"],
        "LAST": params["last_sprint"],
        "PENCILS_DOWN": params.get("pencils_down", params["last_sprint"]),
        "remaining_sprint_count": params.get("remaining_sprints", "?"),
        "total_dev_sprints": params.get("total_dev_sprints", "?"),
        "component_filter": meta.get("component_filter", "none"),
        "assessed_features": meta.get("assessed_features", 0),
        "total_features": meta.get("total_features", 0),
        "data_quality_failures": meta.get("data_quality_failures", 0),
        "overall_risk": f"**{meta.get('overall_risk', 'LOW')}**",
    }

    for key, val in header_map.items():
        template = template.replace(f"{{{key}}}", str(val))

    sections = {}

    # Data Quality
    dq_rows = [[d["feature_key"], str(d["epic_count"]), str(d["story_count"]),
                 f"{d['pointed_pct']}%", d["status"]] for d in checks["data_quality"]]
    dq_table = render_md_table(["Feature", "Epics", "Stories", "Pointed %", "Status"], dq_rows)
    fail_items = [d for d in checks["data_quality"] if d["status"] == "FAIL"]
    fail_text = ""
    if fail_items:
        fail_text = "\n\n### Features Needing Breakdown\n\n" + "\n".join(
            f"- **{d['feature_key']}** — {d['reason']}" for d in fail_items
        )
    sections["DATA_QUALITY"] = f"## Data Quality\n\n{dq_table}{fail_text}"

    # Capacity
    cap_rows = [[c["display_name"], str(c["assigned_sp"]), str(c["remaining_capacity"]),
                  str(c["overrun"]), c["status"], ", ".join(c["features"])] for c in checks["capacity"]]
    cap_table = render_md_table(["Person", "Assigned SP", "Remaining Capacity", "Overrun", "Status", "Features"], cap_rows)
    cap_recs = recs.get("per_person", [])
    cap_rec_text = ""
    if cap_recs:
        cap_rec_text = "\n\n### Recommendations\n\n" + "\n".join(f"- {r}" for r in cap_recs)
    sections["CAPACITY"] = f"## Capacity by Person\n\n{cap_table}{cap_rec_text}"

    # Timeline
    tl_rows = [[t["feature_key"], str(t["remaining_sp"]), str(t["velocity_per_sprint"]),
                 str(t["sprints_needed"]), str(t["sprints_left"]), str(t["gap"]), t["risk"]]
                for t in checks["timeline"] if t["risk"] != "N/A"]
    tl_table = render_md_table(["Feature", "Remaining SP", "Velocity/Sprint", "Sprints Needed", "Sprints Left", "Gap", "Risk"], tl_rows)
    feat_recs = recs.get("per_feature", [])
    feat_rec_text = ""
    if feat_recs:
        feat_rec_text = "\n\n### Recommendations\n\n" + "\n".join(f"- {r}" for r in feat_recs)
    sections["TIMELINE"] = f"## Feature Timeline Risk\n\n{tl_table}{feat_rec_text}"

    # Assignment
    spof_rows = [[s["feature_key"], s["sole_contributor_display"], ""]
                  for s in checks["assignment"]["spof"]]
    spof_table = render_md_table(["Feature", "Sole Contributor", "Action"], spof_rows)
    unassigned_rows = [[u["feature_key"], str(u["count"]), str(u["sp"]), ""]
                        for u in checks["assignment"]["unassigned"]]
    unassigned_table = render_md_table(["Feature", "Unassigned Stories", "Unassigned SP", "Action"], unassigned_rows)
    sections["ASSIGNMENT"] = f"## Assignment Risks\n\n### Single Points of Failure\n\n{spof_table}\n\n### Unassigned Work\n\n{unassigned_table}"

    # Bug Load
    bc_rows = [[b["key"], b["priority"], b["component"], b["status"], ""]
                for b in checks["bug_load"]["unassigned_blocker_critical"]]
    bc_table = render_md_table(["Bug", "Priority", "Component", "Status", "Action"], bc_rows)
    comp_rows = []
    for comp, data in checks["bug_load"].get("by_component", {}).items():
        comp_rows.append([comp, str(data["total"]), str(data["blocker"]), str(data["critical"]), str(data["unassigned"])])
    comp_table = render_md_table(["Component", "Total", "Blocker", "Critical", "Unassigned"], comp_rows)
    sections["BUG_LOAD"] = f"## Bug Load\n\n### Unassigned Blocker/Critical Bugs\n\n{bc_table}\n\n### Bug Summary by Component\n\n{comp_table}"

    # Sizing
    sz_rows = [[s["feature_key"], s["tshirt"], str(s["actual_sp"]), str(s["epic_count"]),
                 str(s["contributor_count"]), s["assessment"]] for s in checks["sizing"]]
    sz_table = render_md_table(["Feature", "T-Shirt Size", "Actual SP", "Epics", "Contributors", "Assessment"], sz_rows)
    sections["SIZING"] = f"## Sizing Mismatches\n\n{sz_table}"

    # Composite
    comp_risk_rows = [[c["feature_key"], c["data_quality"], c["timeline"], c["capacity"],
                        c["assignment"], c["bugs"], c["sizing"], f"**{c['composite_risk']}**"]
                       for c in checks["composite"]]
    comp_risk_table = render_md_table(
        ["Feature", "Data Quality", "Timeline", "Capacity", "Assignment", "Bugs", "Sizing", "Composite Risk"],
        comp_risk_rows
    )
    sections["PROGRESS"] = f"## Composite Risk Assessment\n\n{comp_risk_table}"

    # Recommendations
    exec_summary = recs.get("executive_summary", "No recommendations generated.")
    team_recs = recs.get("team_level", [])
    rec_parts = ["## Recommendations"]
    if recs.get("per_person"):
        rec_parts.append("\n### Per-Person Actions\n\n" + "\n".join(f"- {r}" for r in recs["per_person"]))
    if recs.get("per_feature"):
        rec_parts.append("\n### Per-Feature Actions\n\n" + "\n".join(f"- {r}" for r in recs["per_feature"]))
    if team_recs:
        rec_parts.append("\n### Team-Level Actions\n\n" + "\n".join(f"- {r}" for r in team_recs))
    sections["RECOMMENDATIONS"] = "\n".join(rec_parts)

    template = template.replace("{executive_recommendation}", exec_summary)
    for key, content in sections.items():
        template = template.replace(f"{{{key}}}", content)

    return jira_linkify(template)


# --- DOCX Rendering ---


def render_docx(checks, recs, params, output_path):
    """Render DOCX directly from structured data."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
    HEADING_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
    HIGH_RED_BG, HIGH_RED_TEXT = "FADBD8", RGBColor(0xC0, 0x39, 0x2B)
    MED_YELLOW_BG, MED_YELLOW_TEXT = "FEF9E7", RGBColor(0xB7, 0x95, 0x0B)
    LOW_GREEN_BG, LOW_GREEN_TEXT = "D5F5E3", RGBColor(0x1E, 0x8E, 0x3E)
    ALT_ROW_BG = "F2F2F2"

    RISK_COLORS = {
        "HIGH": (HIGH_RED_BG, HIGH_RED_TEXT),
        "OVER": (HIGH_RED_BG, HIGH_RED_TEXT),
        "FAIL": (HIGH_RED_BG, HIGH_RED_TEXT),
        "MEDIUM": (MED_YELLOW_BG, MED_YELLOW_TEXT),
        "WARN": (MED_YELLOW_BG, MED_YELLOW_TEXT),
        "LOW": (LOW_GREEN_BG, LOW_GREEN_TEXT),
        "OK": (LOW_GREEN_BG, LOW_GREEN_TEXT),
        "PASS": (LOW_GREEN_BG, LOW_GREEN_TEXT),
    }

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    for level, sz, color in [("Heading 1", 20, DARK_BLUE), ("Heading 2", 16, HEADING_BLUE), ("Heading 3", 13, HEADING_BLUE)]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(sz)
        s.font.color.rgb = color
        s.font.bold = True
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def set_shading(cell, color_hex):
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)

    def add_styled_table(headers, rows, risk_col_indices=None):
        if risk_col_indices is None:
            risk_col_indices = set()
        table = doc.add_table(rows=1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = True

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_shading(cell, "2C5F8A")

        for ri, row_data in enumerate(rows):
            row_cells = table.add_row().cells
            for ci, val in enumerate(row_data):
                if ci >= len(headers):
                    break
                cell = row_cells[ci]
                cell.text = ""
                text = str(val).replace("**", "")
                run = cell.paragraphs[0].add_run(text)
                run.font.size = Pt(10)

                if ci in risk_col_indices:
                    clean = text.strip().upper()
                    if clean in RISK_COLORS:
                        bg, fg = RISK_COLORS[clean]
                        set_shading(cell, bg)
                        run.font.color.rgb = fg
                        run.bold = True
                elif ri % 2 == 1:
                    set_shading(cell, ALT_ROW_BG)

        tbl_pr = table._tbl.tblPr if table._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
            f'</w:tblBorders>'
        )
        tbl_pr.append(borders)
        doc.add_paragraph()

    def add_bullets(items):
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(item))

    meta = checks["meta"]

    doc.add_heading(f"Release Planning Risk Assessment: OCP {params['version']}", level=1)

    info_headers = ["", ""]
    pencils_down = params.get("pencils_down", params["last_sprint"])
    info_rows = [
        ["Analysis Date", params["today"]],
        ["Release Window", f"Sprint {params['first_sprint']} -- Sprint {params['last_sprint']}"],
        ["Pencils Down", f"Sprint {pencils_down}"],
        ["Sprints Remaining", f"{params.get('remaining_sprints', '?')} of {params.get('total_dev_sprints', '?')} dev sprints (to pencils down)"],
        ["Component Filter", meta.get("component_filter", "none")],
        ["Features Assessed", f"{meta['assessed_features']} / {meta['total_features']}"],
        ["Data Quality Failures", str(meta["data_quality_failures"])],
        ["Overall Risk", meta["overall_risk"]],
    ]
    add_styled_table(info_headers, info_rows, risk_col_indices={1})

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(recs.get("executive_summary", ""))

    doc.add_heading("Data Quality", level=2)
    dq_rows = [[d["feature_key"], d["epic_count"], d["story_count"], f"{d['pointed_pct']}%", d["status"]]
                for d in checks["data_quality"]]
    add_styled_table(["Feature", "Epics", "Stories", "Pointed %", "Status"], dq_rows, risk_col_indices={4})

    doc.add_heading("Capacity by Person", level=2)
    cap_rows = [[c["display_name"], c["assigned_sp"], c["remaining_capacity"], c["overrun"], c["status"],
                  ", ".join(c["features"])] for c in checks["capacity"]]
    add_styled_table(["Person", "Assigned SP", "Remaining Capacity", "Overrun", "Status", "Features"], cap_rows, risk_col_indices={4})

    doc.add_heading("Feature Timeline Risk", level=2)
    tl_rows = [[t["feature_key"], t["remaining_sp"], t["velocity_per_sprint"], t["sprints_needed"],
                 t["sprints_left"], t["gap"], t["risk"]]
                for t in checks["timeline"] if t["risk"] != "N/A"]
    add_styled_table(["Feature", "Remaining SP", "Velocity/Sprint", "Sprints Needed", "Sprints Left", "Gap", "Risk"], tl_rows, risk_col_indices={6})

    doc.add_heading("Assignment Risks", level=2)
    doc.add_heading("Single Points of Failure", level=3)
    spof_rows = [[s["feature_key"], s["sole_contributor_display"]] for s in checks["assignment"]["spof"]]
    add_styled_table(["Feature", "Sole Contributor"], spof_rows)
    doc.add_heading("Unassigned Work", level=3)
    ua_rows = [[u["feature_key"], u["count"], u["sp"]] for u in checks["assignment"]["unassigned"]]
    add_styled_table(["Feature", "Unassigned Stories", "Unassigned SP"], ua_rows)

    doc.add_heading("Bug Load", level=2)
    bc_rows = [[b["key"], b["priority"], b["component"], b["status"]]
                for b in checks["bug_load"]["unassigned_blocker_critical"]]
    add_styled_table(["Bug", "Priority", "Component", "Status"], bc_rows)

    doc.add_heading("Sizing Mismatches", level=2)
    sz_rows = [[s["feature_key"], s["tshirt"], s["actual_sp"], s["epic_count"],
                 s["contributor_count"], s["assessment"]] for s in checks["sizing"]]
    add_styled_table(["Feature", "T-Shirt", "Actual SP", "Epics", "Contributors", "Assessment"], sz_rows, risk_col_indices={5})

    doc.add_heading("Composite Risk Assessment", level=2)
    cr_rows = [[c["feature_key"], c["data_quality"], c["timeline"], c["capacity"],
                 c["assignment"], c["bugs"], c["sizing"], c["composite_risk"]]
                for c in checks["composite"]]
    add_styled_table(["Feature", "Data Quality", "Timeline", "Capacity", "Assignment", "Bugs", "Sizing", "Composite Risk"],
                     cr_rows, risk_col_indices={1, 2, 3, 4, 5, 6, 7})

    doc.add_heading("Recommendations", level=2)
    if recs.get("per_person"):
        doc.add_heading("Per-Person Actions", level=3)
        add_bullets(recs["per_person"])
    if recs.get("per_feature"):
        doc.add_heading("Per-Feature Actions", level=3)
        add_bullets(recs["per_feature"])
    if recs.get("team_level"):
        doc.add_heading("Team-Level Actions", level=3)
        add_bullets(recs["team_level"])

    doc.save(output_path)
    print(f"Wrote {output_path}", file=sys.stderr)


# --- Main ---


def main():
    parser = argparse.ArgumentParser(description="Assemble release planning report")
    parser.add_argument("--checks", required=True, help="Path to checks.json")
    parser.add_argument("--recommendations", required=True, help="Path to recommendations.json")
    parser.add_argument("--template", required=True, help="Path to report template .md")
    parser.add_argument("--version", required=True)
    parser.add_argument("--today", required=True)
    parser.add_argument("--first-sprint", required=True)
    parser.add_argument("--last-sprint", required=True)
    parser.add_argument("--pencils-down", default=None, help="Pencils down sprint (defaults to last sprint)")
    parser.add_argument("--remaining-sprints", required=True)
    parser.add_argument("--total-dev-sprints", required=True)
    parser.add_argument("--output", required=True, help="Output base path (without extension)")
    args = parser.parse_args()

    checks = load_json(args.checks)
    recs = load_json(args.recommendations)
    with open(args.template, "r") as f:
        template = f.read()

    params = {
        "version": args.version,
        "today": args.today,
        "first_sprint": args.first_sprint,
        "last_sprint": args.last_sprint,
        "pencils_down": args.pencils_down or args.last_sprint,
        "remaining_sprints": args.remaining_sprints,
        "total_dev_sprints": args.total_dev_sprints,
    }

    md_output = args.output + ".md"
    docx_output = args.output + ".docx"

    md_content = render_markdown(checks, recs, template, params)
    parent = os.path.dirname(md_output)
    if parent:
        os.makedirs(parent, exist_ok=True)
    with open(md_output, "w") as f:
        f.write(md_content)
    print(f"Wrote {md_output}", file=sys.stderr)

    render_docx(checks, recs, params, docx_output)


if __name__ == "__main__":
    main()
