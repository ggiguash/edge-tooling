#!/usr/bin/env python3
"""Convert a markdown release-planning report to a styled DOCX file."""

import argparse
import os
import re
import sys

from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# --- Colors ---

DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
HEADING_BLUE = RGBColor(0x2C, 0x5F, 0x8A)
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)
TABLE_HEADER_BG = "D6E4F0"
TABLE_ALT_ROW_BG = "F2F2F2"
HIGH_RED_BG = "FADBD8"
MEDIUM_YELLOW_BG = "FEF9E7"
LOW_GREEN_BG = "D5F5E3"
HIGH_RED_TEXT = RGBColor(0xC0, 0x39, 0x2B)
MEDIUM_YELLOW_TEXT = RGBColor(0xB7, 0x95, 0x0B)
LOW_GREEN_TEXT = RGBColor(0x1E, 0x8E, 0x3E)


# --- Markdown Parsing ---


def parse_md_lines(text):
    """Split markdown into semantic blocks."""
    lines = text.split("\n")
    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
            i += 1
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
            i += 1
        elif line.startswith("| ") or line.startswith("|"):
            table_lines = []
            while i < len(lines) and (lines[i].startswith("|") or lines[i].strip().startswith("|")):
                table_lines.append(lines[i])
                i += 1
            blocks.append(("table", table_lines))
        elif line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:].strip())
                i += 1
            blocks.append(("bullet_list", items))
        elif line.startswith("> "):
            blocks.append(("blockquote", line[2:].strip()))
            i += 1
        elif line.strip() == "---":
            blocks.append(("separator", ""))
            i += 1
        elif line.strip() == "":
            i += 1
        else:
            blocks.append(("paragraph", line.strip()))
            i += 1

    return blocks


def parse_table(table_lines):
    """Parse markdown table lines into header + rows."""
    if len(table_lines) < 2:
        return [], []

    def split_row(line, col_count=None):
        stripped = line.strip().strip("|")
        cells = re.split(r'(?<!\\)\|', stripped)
        cells = [c.strip().replace("\\|", "|") for c in cells]
        if col_count and len(cells) > col_count:
            cells = cells[:col_count - 1] + [" | ".join(cells[col_count - 1:])]
        return cells

    header = split_row(table_lines[0])
    col_count = len(header)

    rows = []
    for line in table_lines[1:]:
        stripped = line.strip()
        if stripped and not re.match(r"^\|[-:\s|]+\|$", stripped):
            rows.append(split_row(line, col_count))

    return header, rows


INLINE_RE = re.compile(
    r"(\*\*(.+?)\*\*)"
    r"|(\*(.+?)\*)"
    r"|(\[([^\]]+)\]\(([^)]+)\))"
)


def parse_inline(text):
    """Parse inline markdown formatting into segments."""
    segments = []
    last_end = 0

    for m in INLINE_RE.finditer(text):
        if m.start() > last_end:
            segments.append(("text", text[last_end:m.start()]))

        if m.group(2):
            segments.append(("bold", m.group(2)))
        elif m.group(4):
            segments.append(("italic", m.group(4)))
        elif m.group(6):
            segments.append(("link", m.group(6), m.group(7)))

        last_end = m.end()

    if last_end < len(text):
        segments.append(("text", text[last_end:]))

    return segments if segments else [("text", text)]


def get_risk_level(text):
    """Detect risk level from cell text."""
    upper = text.upper().strip()
    if upper in ("HIGH", "🔴", "FAIL"):
        return "high"
    if upper in ("MEDIUM", "🟡", "WARN"):
        return "medium"
    if upper in ("LOW", "🟢", "PASS"):
        return "low"
    return None


# --- DOCX Building ---


def add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" {nsdecls("r")}/>')
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="Hyperlink"/>'
        f'<w:color w:val="{str(LINK_BLUE)}"/>'
        f'<w:u w:val="single"/>'
        f'</w:rPr>'
        f'<w:t xml:space="preserve">{xml_escape(text)}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._element.append(hyperlink)


def add_formatted_text(paragraph, text, font_size=None):
    """Add text with inline formatting (bold, italic, links) to a paragraph."""
    segments = parse_inline(text)
    for seg in segments:
        if seg[0] == "bold":
            run = paragraph.add_run(seg[1])
            run.bold = True
            if font_size:
                run.font.size = font_size
        elif seg[0] == "italic":
            run = paragraph.add_run(seg[1])
            run.italic = True
            if font_size:
                run.font.size = font_size
        elif seg[0] == "link":
            add_hyperlink(paragraph, seg[2], seg[1])
        else:
            run = paragraph.add_run(seg[1])
            if font_size:
                run.font.size = font_size


def set_cell_shading(cell, color_hex):
    """Set background color on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)


def style_document(doc):
    """Configure base document styles."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    for level, size, color in [("Heading 1", 20, DARK_BLUE), ("Heading 2", 16, HEADING_BLUE), ("Heading 3", 13, HEADING_BLUE)]:
        s = doc.styles[level]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def add_table(doc, header, rows):
    """Add a styled table to the document."""
    if not header:
        return

    col_count = len(header)
    table = doc.add_table(rows=1, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        if i < col_count:
            hdr_cells[i].text = ""
            p = hdr_cells[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_shading(hdr_cells[i], "2C5F8A")

    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            if i < col_count:
                row_cells[i].text = ""
                p = row_cells[i].paragraphs[0]
                add_formatted_text(p, text, font_size=Pt(10))

                risk = get_risk_level(text)
                if risk == "high":
                    set_cell_shading(row_cells[i], HIGH_RED_BG)
                    for run in p.runs:
                        run.font.color.rgb = HIGH_RED_TEXT
                        run.bold = True
                elif risk == "medium":
                    set_cell_shading(row_cells[i], MEDIUM_YELLOW_BG)
                    for run in p.runs:
                        run.font.color.rgb = MEDIUM_YELLOW_TEXT
                        run.bold = True
                elif risk == "low":
                    set_cell_shading(row_cells[i], LOW_GREEN_BG)
                    for run in p.runs:
                        run.font.color.rgb = LOW_GREEN_TEXT

                if not risk and row_idx % 2 == 1:
                    set_cell_shading(row_cells[i], TABLE_ALT_ROW_BG)

    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="BFBFBF"/>'
        f'</w:tblBorders>'
    )
    tbl_pr.append(borders)


def convert(md_path, docx_path):
    """Convert a markdown report to DOCX."""
    with open(md_path, "r") as f:
        text = f.read()

    blocks = parse_md_lines(text)
    doc = Document()
    style_document(doc)

    for block in blocks:
        kind = block[0]
        content = block[1]

        if kind == "h1":
            doc.add_heading(content, level=1)
        elif kind == "h2":
            doc.add_heading(content, level=2)
        elif kind == "h3":
            doc.add_heading(content, level=3)
        elif kind == "table":
            header, rows = parse_table(content)
            add_table(doc, header, rows)
            doc.add_paragraph()
        elif kind == "bullet_list":
            for item in content:
                p = doc.add_paragraph(style="List Bullet")
                p.text = ""
                add_formatted_text(p, item)
        elif kind == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.27)
            run = p.add_run(content)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif kind == "separator":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="D9D9D9"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            add_formatted_text(p, content)

    doc.save(docx_path)
    print(f"Wrote {docx_path}", file=sys.stderr)


def main():
    parser = argparse.ArgumentParser(description="Convert markdown report to styled DOCX")
    parser.add_argument("input", help="Input markdown file path")
    parser.add_argument("--output", "-o", help="Output DOCX file path (default: same name with .docx extension)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: {args.input} not found", file=sys.stderr)
        sys.exit(1)

    output = args.output or os.path.splitext(args.input)[0] + ".docx"
    convert(args.input, output)


if __name__ == "__main__":
    main()
